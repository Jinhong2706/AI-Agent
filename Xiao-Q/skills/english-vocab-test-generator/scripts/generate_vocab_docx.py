#!/usr/bin/env python3
"""
英语词汇试卷生成器 - MaryMa
支持两种模式：
1. 单词模式（word）：中文意思 + 词性 + 括号（右对齐）
   - 输入格式：英文 音标 词性 中文  或  英文 词性 中文
2. 词组模式（vocab）：中文意思 + 括号（右对齐，无英文、无词性）
   - 输入格式：每行 = 中文意思（英文答案仅用于生成参考答案，不出现在试卷上）

更新：单列布局，括号靠右对齐。词组模式下不显示英文，学生根据中文写英文。
"""

import sys
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn


def generate_vocab_test(words_data, output_path, title="英语词汇测试", mode="word"):
    """
    生成词汇试卷 Word 文档
    
    Args:
        words_data: 单词数据列表，每项为 dict，包含 chinese, pos, english, phonetic
        output_path: 输出文件路径
        title: 文档标题
        mode: 模式，"word"=单词模式（含词性），"vocab"=词汇模式（无词性）
    """
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
    
    # 添加标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.name = '宋体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title_run.font.size = Pt(15)  # 小三号
    title_run.font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加空行
    doc.add_paragraph()
    
    # 创建表格（单列：中文意思+词性+括号）
    table = doc.add_table(rows=0, cols=1)  # 不创建初始行
    table.style = 'Table Grid'
    
    # 设置表格宽度
    table.autofit = False
    table.allow_autofit = False
    
    # 获取页面宽度（A4纸约21cm，减去边距）
    page_width = Cm(21)
    left_margin = Cm(3.18)
    right_margin = Cm(3.18)
    available_width = page_width - left_margin - right_margin
    
    # 不添加表头（按用户要求删除）
    
    # 添加单词数据
    for word in words_data:
        row_cells = table.add_row().cells
        
        # 获取数据
        chinese = word.get('chinese', '')
        pos = word.get('pos', '')
        
        # 构建左列文本
        if mode == "vocab":
            # 词汇模式：中文意思  (        )
            left_text = f"{chinese}  (        )"
        else:
            # 单词模式：中文意思  词性  (        )
            if pos:
                left_text = f"{chinese}  {pos}  (        )"
            else:
                left_text = f"{chinese}  (        )"
        
        # 设置单元格内容 - 使用制表符将括号推到右侧
        cell = row_cells[0]
        
        # 清除单元格内容
        cell.text = ''
        
        for paragraph in cell.paragraphs:
            # 添加内容：中文(+词性) + 空格 + 左括号 + 制表符 + 右括号
            if mode == "vocab" or not pos:
                # 词汇模式：中文 + 空格 + 左括号 + 制表符 + 右括号
                run1 = paragraph.add_run(chinese)
                run1.font.name = '宋体'
                run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run1.font.size = Pt(12)
                
                # 添加空格 + 左括号
                run_left = paragraph.add_run(' (')
                run_left.font.name = '宋体'
                run_left._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run_left.font.size = Pt(12)
                
                # 添加制表符（将右括号推到最右侧）
                run_tab = paragraph.add_run('\t')
                
                # 右括号
                run_right = paragraph.add_run(')')
                run_right.font.name = '宋体'
                run_right._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run_right.font.size = Pt(12)
            else:
                # 单词模式：中文 + 词性 + 空格 + 左括号 + 制表符 + 右括号
                run1 = paragraph.add_run(chinese)
                run1.font.name = '宋体'
                run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run1.font.size = Pt(12)
                
                # 词性部分
                run_pos = paragraph.add_run(f'  {pos}')
                run_pos.font.name = '宋体'
                run_pos._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run_pos.font.size = Pt(12)
                
                # 添加空格 + 左括号
                run_left = paragraph.add_run(' (')
                run_left.font.name = '宋体'
                run_left._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run_left.font.size = Pt(12)
                
                # 添加制表符（将右括号推到最右侧）
                run_tab = paragraph.add_run('\t')
                
                # 右括号
                run_right = paragraph.add_run(')')
                run_right.font.name = '宋体'
                run_right._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run_right.font.size = Pt(12)
            
            # 设置段落格式
            paragraph.paragraph_format.line_spacing = 1.5
    
    # 设置列宽并添加右对齐制表符
    for row in table.rows:
        cell = row.cells[0]
        cell.width = available_width
        
        # 设置右对齐制表符（在单元格最右侧）
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                available_width,
                WD_TAB_ALIGNMENT.RIGHT
            )
    
    # 保存文档
    doc.save(output_path)
    return output_path


def parse_input_text(text, mode="word"):
    """
    解析输入文本，提取单词/词组数据

    单词模式（word）：
      英文 音标 词性 中文
      例：ability /əˈbɪləti/ n. 能力
    词组模式（vocab）：
      每行格式 = 中文意思（英文答案可选，用于生成参考答案）
      例：天气怎么样？
      或：What is the weather like? = How is the weather? 天气怎么样？

    Args:
        text: 输入文本
        mode: 模式，"word"=单词模式，"vocab"=词组模式
    """
    words = []
    lines = text.strip().split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if mode == "vocab":
            # 词组模式：整行作为中文意思（如果包含英文前缀，提取纯中文部分）
            # 智能检测：查找连续中文字符作为中文意思
            import re
            # 尝试找中文部分（从第一个中文字符开始到最后）
            chinese_match = re.search(r'[\u4e00-\u9fff]', line)
            if chinese_match:
                start_idx = chinese_match.start()
                chinese = line[start_idx:].strip()
                english_prefix = line[:start_idx].strip()
                words.append({
                    'english': english_prefix,
                    'phonetic': '',
                    'pos': '',
                    'chinese': chinese
                })
            else:
                # 纯中文行
                words.append({
                    'english': '',
                    'phonetic': '',
                    'pos': '',
                    'chinese': line
                })
        else:
            # 单词模式（含词性）
            if len(parts) >= 4:
                # 格式：英文 音标 词性 中文
                words.append({
                    'english': parts[0],
                    'phonetic': parts[1],
                    'pos': parts[2],
                    'chinese': ' '.join(parts[3:])
                })
            elif len(parts) == 3:
                # 格式：英文 词性 中文（无音标）
                words.append({
                    'english': parts[0],
                    'phonetic': '',
                    'pos': parts[1],
                    'chinese': parts[2]
                })
            elif len(parts) >= 2:
                # 简单格式：英文 中文
                words.append({
                    'english': parts[0],
                    'phonetic': '',
                    'pos': '',
                    'chinese': ' '.join(parts[1:])
                })
    
    return words


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法: python generate_vocab_docx.py <输入文件> <输出文件> [标题] [模式]")
        print("示例: python generate_vocab_docx.py words.txt output.docx '六年级词汇测试' word")
        print("模式: word=单词模式(含词性), vocab=词汇模式(无词性)")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    title = sys.argv[3] if len(sys.argv) > 3 else "英语词汇测试"
    mode = sys.argv[4] if len(sys.argv) > 4 else "word"
    
    # 读取输入文件
    with open(input_file, 'r', encoding='utf-8') as f:
        text = f.read()
    
    # 解析单词数据
    words_data = parse_input_text(text, mode)
    
    if not words_data:
        print("错误：未能从输入文件中解析出单词数据")
        sys.exit(1)
    
    # 生成文档
    result = generate_vocab_test(words_data, output_file, title, mode)
    print(f"文档已生成: {result}")
    print(f"共包含 {len(words_data)} 个单词")
    print(f"模式: {'词汇模式(无词性)' if mode == 'vocab' else '单词模式(含词性)'}")
