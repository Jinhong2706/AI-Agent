#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
艺术作品解读 - Word报告生成器
根据六维/四维框架生成艺术鉴赏报告
"""

import sys
import json
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("请先安装依赖: pip install python-docx")
    sys.exit(1)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = docx.oxml.OxmlElement(tag)
                tcPr.append(element)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), edge_data.get('color', '000000'))


def create_cover(doc, artwork_name, artist, year, date_str):
    """创建封面"""
    # 空行
    for _ in range(6):
        doc.add_paragraph()
    
    # 主标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("艺术作品鉴赏报告")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(139, 69, 19)  # 棕色，艺术感
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 空行
    for _ in range(2):
        doc.add_paragraph()
    
    # 作品名称
    artwork_para = doc.add_paragraph()
    artwork_run = artwork_para.add_run(f"《{artwork_name}》")
    artwork_run.font.size = Pt(22)
    artwork_run.font.bold = True
    artwork_run.font.color.rgb = RGBColor(51, 51, 51)
    artwork_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 空行
    doc.add_paragraph()
    
    # 艺术家
    artist_para = doc.add_paragraph()
    artist_run = artist_para.add_run(f"艺术家：{artist}")
    artist_run.font.size = Pt(14)
    artist_run.font.color.rgb = RGBColor(102, 102, 102)
    artist_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 创作年代
    if year:
        year_para = doc.add_paragraph()
        year_run = year_para.add_run(f"创作年代：{year}")
        year_run.font.size = Pt(14)
        year_run.font.color.rgb = RGBColor(102, 102, 102)
        year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 空行
    for _ in range(4):
        doc.add_paragraph()
    
    # 日期
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(date_str)
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 分页
    doc.add_page_break()


def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(139, 69, 19)  # 棕色
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = RGBColor(160, 82, 45)  # 深棕色
            run.font.size = Pt(14)
        else:
            run.font.color.rgb = RGBColor(51, 51, 51)
    return heading


def add_dimension_section(doc, dimension_num, dimension_name, dimension_key, content, color):
    """添加维度分析章节"""
    # 维度标题
    heading = add_heading_custom(doc, f"第{dimension_num}维：{dimension_name}", level=2)
    
    # 内容
    if content:
        for key, value in content.items():
            if value:
                # 子标题
                para = doc.add_paragraph()
                label = para.add_run(f"{key}：")
                label.font.bold = True
                label.font.color.rgb = RGBColor(int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16))
                
                # 内容
                if isinstance(value, list):
                    for item in value:
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(str(item))
                        p.paragraph_format.line_spacing = 1.5
                else:
                    content_para = doc.add_paragraph(str(value))
                    content_para.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()


def create_info_table(doc, artwork_info):
    """创建作品信息表"""
    add_heading_custom(doc, "作品基本信息", level=2)
    
    # 创建表格
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '项目'
    hdr_cells[1].text = '内容'
    
    # 设置表头样式
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加数据行
    info_items = [
        ("作品名称", artwork_info.get('name', '')),
        ("艺术家", artwork_info.get('artist', '')),
        ("创作年代", artwork_info.get('year', '')),
        ("作品类型", artwork_info.get('type', '')),
        ("尺寸/时长", artwork_info.get('size', '')),
        ("收藏地", artwork_info.get('location', '')),
        ("艺术流派", artwork_info.get('movement', '')),
    ]
    
    for label, value in info_items:
        if value:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(value)
    
    doc.add_paragraph()


def generate_report(data, output_path):
    """生成完整的艺术鉴赏报告"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)
    
    # 提取数据
    artwork_info = data.get('artwork', {})
    artwork_name = artwork_info.get('name', '未知作品')
    artist = artwork_info.get('artist', '未知艺术家')
    year = artwork_info.get('year', '')
    date_str = data.get('date', datetime.now().strftime('%Y年%m月%d日'))
    analysis_mode = data.get('mode', 'six')  # 'six' 或 'four'
    
    # 1. 封面
    create_cover(doc, artwork_name, artist, year, date_str)
    
    # 2. 作品基本信息
    create_info_table(doc, artwork_info)
    
    # 3. 分析框架说明
    add_heading_custom(doc, "分析框架", level=2)
    if analysis_mode == 'six':
        framework_text = "本报告采用六维分析法：形式、内容、技法、情感、思想/主题、背景"
    else:
        framework_text = "本报告采用极简四维法：形式、内容、情感、思想（含背景）"
    
    para = doc.add_paragraph(framework_text)
    para.paragraph_format.line_spacing = 1.5
    doc.add_paragraph()
    
    # 分页
    doc.add_page_break()
    
    # 4. 维度分析
    add_heading_custom(doc, "深度分析", level=1)
    doc.add_paragraph()
    
    if analysis_mode == 'six':
        # 六维分析
        dimensions = [
            ("一", "形式维度", 'form', "#FF6B6B"),
            ("二", "内容维度", 'content', "#4ECDC4"),
            ("三", "技法维度", 'technique', "#45B7D1"),
            ("四", "情感维度", 'emotion', "#96CEB4"),
            ("五", "思想/主题维度", 'theme', "#DDA0DD"),
            ("六", "背景维度", 'context', "#F7DC6F"),
        ]
    else:
        # 四维分析
        dimensions = [
            ("一", "形式维度", 'form', "#FF6B6B"),
            ("二", "内容维度", 'content', "#4ECDC4"),
            ("三", "情感维度", 'emotion', "#96CEB4"),
            ("四", "思想维度（含背景）", 'theme', "#DDA0DD"),
        ]
    
    for dim_num, dim_name, dim_key, color in dimensions:
        dim_content = data.get(dim_key, {})
        if dim_content:
            add_dimension_section(doc, dim_num, dim_name, dim_key, dim_content, color)
    
    # 分页
    doc.add_page_break()
    
    # 5. 综合评述
    add_heading_custom(doc, "综合评述", level=1)
    
    # 艺术价值
    artistic_value = data.get('artistic_value', '')
    if artistic_value:
        add_heading_custom(doc, "艺术价值", level=2)
        para = doc.add_paragraph(artistic_value)
        para.paragraph_format.line_spacing = 1.5
        doc.add_paragraph()
    
    # 历史地位
    historical_position = data.get('historical_position', '')
    if historical_position:
        add_heading_custom(doc, "历史地位", level=2)
        para = doc.add_paragraph(historical_position)
        para.paragraph_format.line_spacing = 1.5
        doc.add_paragraph()
    
    # 个人感悟
    personal_insight = data.get('personal_insight', '')
    if personal_insight:
        add_heading_custom(doc, "个人感悟", level=2)
        para = doc.add_paragraph(personal_insight)
        para.paragraph_format.line_spacing = 1.5
        doc.add_paragraph()
    
    # 分页
    doc.add_page_break()
    
    # 6. 延伸阅读
    add_heading_custom(doc, "延伸阅读", level=1)
    
    # 相关作品
    related_works = data.get('related_works', [])
    if related_works:
        add_heading_custom(doc, "相关作品推荐", level=2)
        for work in related_works:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(work)
            para.paragraph_format.line_spacing = 1.5
        doc.add_paragraph()
    
    # 参考资料
    references = data.get('references', [])
    if references:
        add_heading_custom(doc, "参考资料", level=2)
        for ref in references:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(ref)
            para.paragraph_format.line_spacing = 1.5
    
    # 保存文档
    doc.save(output_path)
    print(f"✅ 艺术鉴赏报告已生成: {output_path}")


def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法: python generate_report.py <json文件> [输出路径]")
        print("示例: python generate_report.py data.json 艺术鉴赏报告.docx")
        sys.exit(1)
    
    json_file = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else "艺术鉴赏报告.docx"
    
    # 读取JSON数据
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # 生成报告
    generate_report(data, output_path)


if __name__ == '__main__':
    main()
