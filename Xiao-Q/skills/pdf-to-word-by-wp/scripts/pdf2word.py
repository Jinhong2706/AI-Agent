#!/usr/bin/env python3
"""
PDF 转 Word 转换器（支持 Aspose.PDF）
https://github.com/aspose-pdf/Aspose.PDF-for-Java

使用方法：
  python pdf2word.py input.pdf [output.docx] [mode]

模式选项：
  flow       — 易编辑模式（默认）
  enhanced   — 增强流模式
  textbox    — 高保真模式

示例：
  python pdf2word.py resume.pdf                    # 输出 resume.docx，Flow 模式
  python pdf2word.py resume.pdf out.docx           # 指定输出文件名
  python pdf2word.py resume.pdf out.docx textbox   # 高保真模式
  python pdf2word.py resume.pdf out.docx enhanced   # 增强流模式

依赖：
  pip install pymupdf python-docx
  或提供 jar（lib/pdftool-core-21.8.jar 已内置处理）
"""

import argparse
import sys
import os
import time
import tempfile
import shutil
from pathlib import Path

VERSION = "1.0.0"

# DocSaveOptions.DocFormat（Aspose.PDF 21.8 内联常量）
FORMAT_DOC  = 0   # .doc
FORMAT_DOCX = 1   # .docx

# DocSaveOptions.RecognitionMode（Aspose.PDF 21.8 内联常量）
MODE_FLOW     = 0   # 易编辑
MODE_ENHANCED = 1   # 增强流
MODE_TEXTBOX  = 2   # 高保真

MODE_NAMES = {
    "flow":     (MODE_FLOW,     "Flow（易编辑）"),
    "enhanced": (MODE_ENHANCED, "EnhancedFlow（增强流）"),
    "textbox":  (MODE_TEXTBOX,  "Textbox（高保真）"),
}

DEFAULT_FORMAT    = FORMAT_DOCX
DEFAULT_MODE      = MODE_FLOW
DEFAULT_PROXIMITY = 2.5
DEFAULT_BULLETS   = True


def build_java_classpath(jar_path: str = None) -> str:
    """构建 Java classpath。优先用户指定 jar，其次 lib/ 下找。"""
    jars = []
    if jar_path and Path(jar_path).exists():
        jars.append(jar_path)

    lib_dir = Path(__file__).parent.parent / "lib"
    if lib_dir.exists():
        for j in lib_dir.glob("*.jar"):
            jars.append(str(j.resolve()))

    if not jars:
        raise FileNotFoundError(
            "未找到 PDF 核心库 jar。\n"
            "请提供 --jar 参数指定路径，或将 jar 放入 lib/ 目录。"
        )
    return os.pathsep.join(jars)


def convert_pdf_java(
    input_pdf: str,
    output_docx: str,
    format: int = FORMAT_DOCX,
    mode: int = DEFAULT_MODE,
    jar_path: str = None,
    proximity: float = DEFAULT_PROXIMITY,
    recognize_bullets: bool = DEFAULT_BULLETS,
    java_opts: str = "-Xmx512m",
) -> float:
    """
    通过 Java + Aspose.PDF 转换 PDF → Word。
    jar 内置处理，无须额外 Crack 步骤。
    如同目录下有 license.xml 则加载（可去除水印）。
    返回转换耗时（毫秒），失败抛出异常。
    """
    start = time.time()
    classpath = build_java_classpath(jar_path)

    java_code = f'''\
import com.aspose.pdf.*;
import java.io.*;

public class Pdf2wordRunner {{
    public static void main(String[] args) throws Exception {{
        // 尝试加载 license.xml（如果有的话，可去除水印）
        try {{
            File lic = new File("license.xml");
            if (lic.exists()) {{
                License l = new License();
                l.setLicense(new FileInputStream(lic));
            }}
        }} catch (Exception e) {{ /* 无 license 继续，jar 内置处理 */ }}

        Document doc = new Document("{_escape_java(input_pdf)}");
        DocSaveOptions opts = new DocSaveOptions();
        opts.setFormat({format});               // 0=DOC, 1=DOCX
        opts.setMode({mode});                   // 0=Flow, 1=EnhancedFlow, 2=Textbox
        opts.setRelativeHorizontalProximity({proximity}f);
        opts.setRecognizeBullets({str(recognize_bullets).lower()});
        doc.save("{_escape_java(output_docx)}", opts);
        doc.close();
    }}
    private static String _escape(String s) {{
        return s.replace("\\\\", "\\\\\\\\").replace("\\"", "\\\\\\"");
    }}
}}
'''

    temp_dir = Path(tempfile.mkdtemp())
    try:
        java_file = temp_dir / "Pdf2wordRunner.java"
        java_file.write_text(java_code, encoding="utf-8")

        compile_res = os.system(
            f'javac -encoding UTF-8 -d "{temp_dir}" -cp "{classpath}" '
            f'"{java_file}" >nul 2>&1'
        )
        if compile_res != 0:
            raise RuntimeError("Java 编译失败，请确认 javac 可用且 jar 路径正确。")

        run_res = os.system(
            f'java {java_opts} -cp "{temp_dir}{os.pathsep}{classpath}" '
            f'Pdf2wordRunner >nul 2>&1'
        )
        if run_res != 0:
            raise RuntimeError(f"转换失败（退出码 {run_res}）。")

        if not Path(output_docx).exists():
            raise RuntimeError("转换完成但输出文件不存在。")
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

    return (time.time() - start) * 1000


def _escape_java(s: str) -> str:
    return s.replace("\\", "\\\\").replace('"', '\\"')


# ──────────────────────────────────────────────
# 备选：纯 Python 方案（pymupdf + python-docx）
# ──────────────────────────────────────────────

def convert_pdf_python(
    input_pdf: str,
    output_docx: str,
    mode: str = "flow",
) -> float:
    """纯 Python 方案，无需 Java/Aspose。"""
    try:
        import fitz
        from docx import Document
    except ImportError:
        raise ImportError(
            "请安装依赖：pip install pymupdf python-docx\n"
            "或使用 Java 方案（需要 javac + lib/pdftool-core-21.8.jar）。"
        )

    start = time.time()
    doc = fitz.open(input_pdf)
    docx = Document()

    for page_num in range(len(doc)):
        page = doc[page_num]
        if page_num > 0:
            docx.add_page_break()
        for block in page.get_text("blocks"):
            text = block[4].strip()
            if text:
                docx.add_paragraph(text)

    doc.close()
    docx.save(output_docx)
    return (time.time() - start) * 1000


# ──────────────────────────────────────────────
# 主入口
# ──────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="PDF 转 Word 转换器（Java / 纯 Python 双方案）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("input",  help="输入 PDF 文件路径")
    parser.add_argument("output", nargs="?", help="输出 Word 文件路径（默认：与 PDF 同名）")
    parser.add_argument("mode",   nargs="?", choices=["flow", "enhanced", "textbox"],
                        default="flow", help="识别模式（默认：flow）")
    parser.add_argument("--jar",       dest="jar",       help="PDF 核心库 jar 路径（默认：lib/*.jar）")
    parser.add_argument("--proximity", dest="proximity", type=float, default=2.5,
                        help="文字水平 proximity（默认：2.5）")
    parser.add_argument("--no-bullets", action="store_true",
                        help="不识别项目符号")
    parser.add_argument("--python",    action="store_true",
                        help="强制使用纯 Python 方案（不需要 Java）")
    parser.add_argument("--java-opts", dest="java_opts", default="-Xmx512m",
                        help="Java 虚拟机选项（默认：-Xmx512m）")
    parser.add_argument("--version", action="version", version=f"%(prog)s {VERSION}")
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    if not input_path.exists():
        print(f"[错误] 文件不存在: {input_path}", file=sys.stderr)
        sys.exit(1)

    output_path = Path(args.output).resolve() if args.output else input_path.with_suffix(".docx")
    mode_int, mode_label = MODE_NAMES.get(args.mode, (MODE_FLOW, "Flow（易编辑）"))

    print("=" * 50)
    print(f"[输入] {input_path}")
    print(f"[输出] {output_path}")
    print(f"[模式] {mode_label}")
    print("=" * 50)

    try:
        if args.python:
            cost = convert_pdf_python(str(input_path), str(output_path), args.mode)
        else:
            cost = convert_pdf_java(
                str(input_path), str(output_path),
                format=FORMAT_DOCX,
                mode=mode_int,
                jar_path=args.jar,
                proximity=args.proximity,
                recognize_bullets=not args.no_bullets,
                java_opts=args.java_opts,
            )
        print(f"[完成] 耗时 {cost:.0f} ms ✓")
        print(f"[文件] {output_path}")
    except ImportError as e:
        print(f"[错误] {e}", file=sys.stderr)
        print("\n提示：使用 --python 模式（需要 pip install pymupdf python-docx）", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"[错误] {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
