#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
extract_requirements.py
从 .docx 概要设计说明书提取需求矩阵与功能描述
用法：python extract_requirements.py <docx_path> [--output <json_path>]
"""
import sys
import json
import argparse
from pathlib import Path
from docx import Document


def extract_paragraphs(doc):
    """提取所有段落文本（含标题级别）"""
    paragraphs = []
    for p in doc.paragraphs:
        if p.text.strip():
            style = p.style.name if p.style else 'Normal'
            level = 0
            if 'Heading' in style:
                try:
                    level = int(style.replace('Heading', '').strip())
                except ValueError:
                    level = 0
            paragraphs.append({
                'level': level,
                'text': p.text.strip(),
                'style': style,
            })
    return paragraphs


def extract_tables(doc):
    """提取所有表格数据"""
    tables = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append({
            'index': ti,
            'rows': rows,
        })
    return tables


def find_requirements_matrix(tables):
    """定位 1.2 节需求矩阵表格"""
    for table in tables:
        if len(table['rows']) < 2:
            continue
        header = table['rows'][0]
        if any('一级需求' in c for c in header) and any('二级需求' in c for c in header):
            return table
    return None


def parse_requirements_matrix(matrix_table):
    """解析需求矩阵为 {一级, 二级, 三级}"""
    if not matrix_table:
        return {'一级需求': '', '二级需求': {}, '三级需求': {}}

    rows = matrix_table['rows']
    header = rows[0]
    level1_col = None
    level2_col = None
    level3_col = None

    for i, c in enumerate(header):
        if '一级需求' in c:
            level1_col = i
        elif '二级需求' in c:
            level2_col = i
        elif '三级需求' in c:
            level3_col = i

    level1 = ''
    secondary_map = {}
    tertiary_map = {}

    for row in rows[1:]:
        if level1_col is not None and level1_col < len(row):
            if row[level1_col].strip():
                level1 = row[level1_col].strip()
        if level2_col is not None and level2_col < len(row) and level3_col is not None and level3_col < len(row):
            secondary = row[level2_col].strip()
            tertiary = row[level3_col].strip()
            if secondary:
                if level1:
                    secondary_map[level1] = secondary
            if tertiary:
                if level2_col is not None and level2_col < len(row):
                    secondary = row[level2_col].strip() or '通用'
                    tertiary_map.setdefault(secondary, []).append(tertiary)

    return {
        '一级需求': level1,
        '二级需求': secondary_map,
        '三级需求': tertiary_map,
    }


def extract_function_descriptions(paragraphs):
    """提取 2.3 节功能描述（按三级功能点分组）"""
    # 简化：找到 2.3 章节并提取其下所有三级功能点
    in_section_2_3 = False
    current_level3 = None
    function_descriptions = {}

    for p in paragraphs:
        text = p['text']
        if '2.3' in text and ('功能实现' in text or '功能设计' in text):
            in_section_2_3 = True
            continue
        if in_section_2_3 and p['level'] <= 1 and '2.4' in text:
            in_section_2_3 = False
            break

        if in_section_2_3 and p['level'] == 3:
            current_level3 = text
            function_descriptions.setdefault(current_level3, [])
        elif in_section_2_3 and current_level3 and p['level'] >= 4:
            function_descriptions[current_level3].append(text)

    return function_descriptions


def main():
    parser = argparse.ArgumentParser(description='从 .docx 提取需求矩阵与功能描述')
    parser.add_argument('docx_path', help='概要设计说明书 .docx 路径')
    parser.add_argument('--output', '-o', help='输出 JSON 路径（可选）')
    args = parser.parse_args()

    docx_path = Path(args.docx_path)
    if not docx_path.exists():
        print(f"错误：文件不存在 - {docx_path}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(docx_path))
    paragraphs = extract_paragraphs(doc)
    tables = extract_tables(doc)

    matrix_table = find_requirements_matrix(tables)
    req_matrix = parse_requirements_matrix(matrix_table)
    func_desc = extract_function_descriptions(paragraphs)

    result = {
        'document': str(docx_path),
        'requirements_matrix': req_matrix,
        'function_descriptions': func_desc,
    }

    output = json.dumps(result, ensure_ascii=False, indent=2)

    if args.output:
        Path(args.output).write_text(output, encoding='utf-8')
        print(f"已写入：{args.output}")
    else:
        print(output)


if __name__ == '__main__':
    main()
