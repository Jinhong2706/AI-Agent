#!/usr/bin/env python3
import random
import os
import argparse
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from generate_practice import (
    gen_add_sub_within_10000,
    gen_division_with_remainder,
    gen_mixed_operations,
    remove_table_borders,
    set_paragraph_spacing,
    set_cell_font,
)


def add_page_break(doc):
    from docx.oxml.ns import qn as _qn
    br = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    p = doc.add_paragraph()
    p._element.append(br)


def add_problems_to_doc(doc, problems_dict, title):
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title_para, 1.0)
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.name = '宋体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    section_configs = [
        ("万以内数的加法和减法（共20题）", problems_dict['add_sub'], 5),
        ("有余数的除法（共25题）", problems_dict['division'], 5),
        ("混合运算（共25题）", problems_dict['mixed'], 5),
    ]

    for idx, (section_title, problems, cols) in enumerate(section_configs, 1):
        section_para = doc.add_paragraph()
        set_paragraph_spacing(section_para, 1.0)
        section_run = section_para.add_run(f"{idx}. {section_title}")
        section_run.font.size = Pt(10.5)
        section_run.font.bold = True
        section_run.font.name = '宋体'
        section_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        rows = (len(problems) + cols - 1) // cols
        prob_table = doc.add_table(rows=rows, cols=cols)
        remove_table_borders(prob_table)

        for i, prob in enumerate(problems):
            row_idx = i // cols
            col_idx = i % cols
            cell = prob_table.rows[row_idx].cells[col_idx]
            cell.text = f"{i + 1}、  {prob}"
            set_cell_font(cell, font_name='宋体', font_size=10.5, line_spacing=1.4)


def generate_merged_docx(count, output_path=None, seed=None):
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    style.paragraph_format.line_spacing = 1.0

    today = datetime.now()

    for set_num in range(1, count + 1):
        set_date = today + timedelta(days=set_num - 1)
        set_date_display = set_date.strftime("%Y-%m-%d")

        if seed is not None:
            random.seed(seed + set_num)

        problems_dict = {
            'add_sub': gen_add_sub_within_10000(20),
            'division': gen_division_with_remainder(25),
            'mixed': gen_mixed_operations(25),
        }

        title = f"北京市二年级下学期口算练习 {set_date_display}"
        add_problems_to_doc(doc, problems_dict, title)

        if set_num < count:
            add_page_break(doc)

    if output_path is None:
        date_str = today.strftime("%Y-%m-%d")
        output_dir = os.getcwd()
        output_path = os.path.join(output_dir, f'oral_calc_practice_merged_{date_str}.docx')

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description='生成多套口算练习题合并到单个Word文件')
    parser.add_argument('--output', '-o', type=str, default=None,
                        help='输出docx文件路径')
    parser.add_argument('--seed', '-s', type=int, default=None,
                        help='随机种子')
    parser.add_argument('--count', '-n', type=int, default=1,
                        help='生成套数（1～20）')
    args = parser.parse_args()

    count = max(1, min(args.count, 20))
    output_path = generate_merged_docx(count, output_path=args.output, seed=args.seed)
    print(f"合并练习题已生成：{output_path}")


if __name__ == '__main__':
    main()
