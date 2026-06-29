#!/usr/bin/env python3
import random
import os
import argparse
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_paragraph_spacing(paragraph, line_spacing=1.5):
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing


def set_cell_font(cell, font_name='宋体', font_size=10.5, bold=False, line_spacing=1.5):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(paragraph, line_spacing)
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def gen_add_sub_within_10000(count=20):
    problems = []
    types = ['four_add', 'four_sub', 'three_add_three', 'three_sub_three']

    attempts = 0
    while len(problems) < count and attempts < count * 20:
        attempts += 1
        t = random.choice(types)

        if t == 'four_add':
            a = random.randint(1000, 9000)
            b = random.randint(100, 9000)
            if a + b > 9999:
                continue
            expr = f"{a} + {b} ="

        elif t == 'four_sub':
            a = random.randint(1000, 9999)
            b = random.randint(100, a)
            expr = f"{a} - {b} ="

        elif t == 'three_add_three':
            a = random.randint(100, 999)
            b = random.randint(100, 999)
            if a + b > 9999:
                continue
            expr = f"{a} + {b} ="

        elif t == 'three_sub_three':
            a = random.randint(100, 999)
            b = random.randint(10, a)
            expr = f"{a} - {b} ="

        else:
            continue

        if expr not in problems:
            problems.append(expr)

    return problems[:count]


def gen_division_with_remainder(count=25):
    problems = []
    for _ in range(count * 5):
        if len(problems) >= count:
            break
        divisor = random.randint(2, 9)
        quotient = random.randint(1, 9)
        remainder = random.randint(1, divisor - 1)
        dividend = divisor * quotient + remainder
        if dividend > 99 or dividend < 5:
            continue
        expr = f"{dividend} ÷ {divisor} ="
        if expr not in problems:
            problems.append(expr)
    while len(problems) < count:
        divisor = random.randint(2, 9)
        quotient = random.randint(1, 9)
        remainder = random.randint(1, divisor - 1)
        dividend = divisor * quotient + remainder
        if dividend < 5 or dividend > 99:
            continue
        expr = f"{dividend} ÷ {divisor} ="
        problems.append(expr)
    return problems[:count]


def gen_mixed_operations(count=25):
    problems = []
    templates = [
        'sub_div', 'add_div', 'sub_div', 'add_div',
        'mul_sub', 'mul_add',
        'div_sub', 'div_add',
    ]
    attempts = 0
    while len(problems) < count and attempts < count * 20:
        attempts += 1
        t = random.choice(templates)

        if t == 'sub_div':
            b = random.randint(2, 9)
            c = random.randint(2, 9)
            a = b * c
            d = random.randint(10, 99)
            result = d - a // b
            if result <= 0:
                continue
            expr = f"{d} - {a} ÷ {b} ="

        elif t == 'add_div':
            b = random.randint(2, 9)
            c = random.randint(2, 9)
            a = b * c
            d = random.randint(10, 99)
            result = d + a // b
            expr = f"{d} + {a} ÷ {b} ="

        elif t == 'mul_sub':
            a = random.randint(2, 9)
            b = random.randint(2, 9)
            product = a * b
            d = random.randint(product + 1, 99)
            result = d - a * b
            if result <= 0:
                continue
            expr = f"{d} - {a} × {b} ="

        elif t == 'mul_add':
            a = random.randint(2, 9)
            b = random.randint(2, 9)
            d = random.randint(10, 99)
            result = d + a * b
            expr = f"{d} + {a} × {b} ="

        elif t == 'div_sub':
            b = random.randint(2, 9)
            c = random.randint(2, 9)
            a = b * c
            d = random.randint(1, a // b - 1)
            if d <= 0:
                continue
            result = a // b - d
            if result <= 0:
                continue
            expr = f"{a} ÷ {b} - {d} ="

        elif t == 'div_add':
            b = random.randint(2, 9)
            c = random.randint(2, 9)
            a = b * c
            d = random.randint(1, 20)
            result = a // b + d
            expr = f"{a} ÷ {b} + {d} ="

        else:
            continue

        if expr not in problems:
            problems.append(expr)

    while len(problems) < count:
        b = random.randint(2, 9)
        c = random.randint(2, 9)
        a = b * c
        d = random.randint(10, 99)
        expr = f"{d} + {a} ÷ {b} ="
        problems.append(expr)

    return problems[:count]


def generate_markdown(problems_dict, title="北京市二年级下学期口算练习"):
    section_configs = [
        ("万以内数的加法和减法（共20题）", problems_dict['add_sub'], 5),
        ("有余数的除法（共25题）", problems_dict['division'], 5),
        ("混合运算（共25题）", problems_dict['mixed'], 5),
    ]

    lines = []
    lines.append(f"# {title}")
    lines.append("")

    for idx, (section_title, problems, cols) in enumerate(section_configs, 1):
        lines.append(f"## {idx}. {section_title}")
        lines.append("")

        header = "|" + "|".join([f" 第{i+1}列 " for i in range(cols)]) + "|"
        separator = "|" + "|".join([" --- " for _ in range(cols)]) + "|"
        lines.append(header)
        lines.append(separator)

        for row_idx in range(0, len(problems), cols):
            row_problems = problems[row_idx:row_idx + cols]
            cells = []
            for col_idx, prob in enumerate(row_problems):
                num = row_idx + col_idx + 1
                cells.append(f" {num}、 {prob} ")
            while len(cells) < cols:
                cells.append(" ")
            lines.append("|" + "|".join(cells) + "|")

        lines.append("")

    return "\n".join(lines)


def generate_docx(problems_dict, title="北京市二年级下学期口算练习", output_path="oral_calc_practice.docx"):
    doc = Document()

    # A4 横版
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)

    # 页边距：上下0.8cm，左右1.5cm，压缩以容纳单页
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # 设置默认段落行间距为1倍
    style = doc.styles['Normal']
    style.paragraph_format.line_spacing = 1.0

    # 标题（宋体，14号，居中，加粗）
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title_para, 1.0)
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.name = '宋体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    section_configs = [
        ("万以内数的加法和减法（共20题）", problems_dict['add_sub'], 5),
        ("有余数的除法（共25题）", problems_dict['division'], 5),
        ("混合运算（共25题）", problems_dict['mixed'], 5),
    ]

    for idx, (section_title, problems, cols) in enumerate(section_configs, 1):
        # 小节标题（宋体，10.5号，加粗）
        section_para = doc.add_paragraph()
        set_paragraph_spacing(section_para, 1.0)
        section_run = section_para.add_run(f"{idx}. {section_title}")
        section_run.font.size = Pt(10.5)
        section_run.font.bold = True
        section_run.font.name = '宋体'
        section_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 题目表格（无边框）
        rows = (len(problems) + cols - 1) // cols
        prob_table = doc.add_table(rows=rows, cols=cols)
        remove_table_borders(prob_table)

        for i, prob in enumerate(problems):
            row_idx = i // cols
            col_idx = i % cols
            cell = prob_table.rows[row_idx].cells[col_idx]
            cell.text = f"{i + 1}、  {prob}"
            set_cell_font(cell, font_name='宋体', font_size=10.5, line_spacing=1.4)

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description='生成北京市二年级下学期口算练习题')
    parser.add_argument('--output', '-o', type=str, default=None,
                        help='输出docx文件路径（默认为当前目录下 oral_calc_practice_年-月-日.docx）')
    parser.add_argument('--seed', '-s', type=int, default=None,
                        help='随机种子，用于生成不同的题目')
    parser.add_argument('--title', '-t', type=str, default=None,
                        help='练习题标题（默认为"北京市二年级下学期口算练习 年-月-日"）')
    parser.add_argument('--count', '-n', type=int, default=1,
                        help='生成套数（1～20）')
    parser.add_argument('--format', '-f', type=str, choices=['docx', 'md'], default='docx',
                        help='输出格式：docx 或 md')
    args = parser.parse_args()

    count = max(1, min(args.count, 20))

    today = datetime.now()
    date_str_display = today.strftime("%Y-%m-%d")
    date_str_file = today.strftime("%Y-%m-%d")

    if args.output:
        output_path = args.output
    else:
        output_dir = os.getcwd()
        ext = 'md' if args.format == 'md' else 'docx'
        output_path = os.path.join(output_dir, f'oral_calc_practice_{date_str_file}.{ext}')

    generated_files = []

    for set_num in range(1, count + 1):
        set_date = today + timedelta(days=set_num - 1)
        set_date_display = set_date.strftime("%Y-%m-%d")
        set_date_file = set_date.strftime("%Y-%m-%d")

        if args.seed is not None:
            random.seed(args.seed + set_num)

        problems_dict = {
            'add_sub': gen_add_sub_within_10000(20),
            'division': gen_division_with_remainder(25),
            'mixed': gen_mixed_operations(25),
        }

        if count == 1:
            file_path = output_path
        else:
            base, ext = os.path.splitext(output_path)
            file_path = f"{base.replace(date_str_file, set_date_file)}{ext}"

        title = args.title if args.title else f"北京市二年级下学期口算练习 {set_date_display}"

        if args.format == 'md':
            md = generate_markdown(problems_dict, title=title)
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(md)
        else:
            generate_docx(problems_dict, title=title, output_path=file_path)

        generated_files.append(file_path)

    for f in generated_files:
        print(f"练习题已生成：{f}")


if __name__ == '__main__':
    main()
