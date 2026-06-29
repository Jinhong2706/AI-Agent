from docx import Document
from docx.oxml.ns import qn
import sys

def convert_docx_to_md(docx_path, md_path):
    doc = Document(docx_path)
    md_lines = []

    def format_runs(paragraph):
        if not paragraph.runs:
            return paragraph.text.strip()
        result = []
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            result.append(text)
        return ''.join(result).strip()

    def process_paragraph(paragraph):
        style_name = paragraph.style.name if paragraph.style else ''
        style_lower = style_name.lower()

        if style_lower.startswith('toc'):
            return ''

        text = format_runs(paragraph)
        if not text:
            return ''

        if style_lower.startswith('heading'):
            level = 1
            for ch in style_lower:
                if ch.isdigit():
                    level = int(ch)
                    break
            return f"{'#' * level} {text}"

        numPr = paragraph._element.find(qn('w:numPr'))
        if numPr is not None:
            ilvl = numPr.find(qn('w:ilvl'))
            level = int(ilvl.get(qn('w:val'))) if ilvl is not None else 0
            indent = '  ' * level
            return f"{indent}- {text}"

        return text

    def process_table(table):
        lines = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            line = '| ' + ' | '.join(cells) + ' |'
            lines.append(line)
            if i == 0:
                lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
        return '\n'.join(lines)

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            for para in doc.paragraphs:
                if para._element == element:
                    result = process_paragraph(para)
                    if result:
                        md_lines.append(result)
                        md_lines.append('')
                    break
        elif tag == 'tbl':
            for table in doc.tables:
                if table._element == element:
                    result = process_table(table)
                    if result:
                        md_lines.append(result)
                        md_lines.append('')
                    break

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md_lines))

    print(f"Converted {docx_path} -> {md_path}")
    print(f"  Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")

if __name__ == '__main__':
    docx_file = sys.argv[1] if len(sys.argv) > 1 else 'input.docx'
    md_file = sys.argv[2] if len(sys.argv) > 2 else 'output.md'
    convert_docx_to_md(docx_file, md_file)
