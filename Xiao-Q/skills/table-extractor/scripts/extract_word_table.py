#!/usr/bin/env python3
"""
Word Table Extractor
Extracts tables from Word (.docx) files and saves to Excel using python-docx.
"""

import sys
import os

def extract_word_table(docx_path, excel_path):
    """Extract tables from Word document and save to Excel"""
    try:
        from docx import Document
        import pandas as pd
        
        # Check if input file exists
        if not os.path.exists(docx_path):
            print(f"Error: Input file '{docx_path}' not found.")
            return False
        
        # Open Word document
        doc = Document(docx_path)
        
        if len(doc.tables) == 0:
            print("No tables found in the Word document.")
            return False
        
        # Extract tables
        tables = []
        for i, table in enumerate(doc.tables):
            # Extract table data
            data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                data.append(row_data)
            
            if data:
                # First row as header
                df = pd.DataFrame(data[1:], columns=data[0] if len(data) > 1 else None)
                tables.append({
                    'table': i + 1,
                    'data': df
                })
        
        if not tables:
            print("No valid tables found in the Word document.")
            return False
        
        # Save to Excel
        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            for idx, table_info in enumerate(tables):
                sheet_name = f"Table_{idx + 1}" if len(tables) > 1 else "Sheet1"
                table_info['data'].to_excel(writer, sheet_name=sheet_name, index=False)
        
        print(f"Successfully extracted {len(tables)} table(s) from '{docx_path}' to '{excel_path}'")
        for table_info in tables:
            print(f"  - Table {table_info['table']}: {len(table_info['data'])} rows")
        return True
        
    except ImportError as e:
        if 'docx' in str(e):
            print("Error: python-docx library not installed.")
            print("Please install it with: pip install python-docx")
        elif 'pandas' in str(e) or 'openpyxl' in str(e):
            print("Error: pandas or openpyxl library not installed.")
            print("Please install it with: pip install pandas openpyxl")
        else:
            print(f"Error: {e}")
        return False
    except Exception as e:
        print(f"Error during extraction: {e}")
        return False

def main():
    if len(sys.argv) != 3:
        print("Usage: python3 extract_word_table.py <input.docx> <output.xlsx>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    excel_path = sys.argv[2]
    
    success = extract_word_table(docx_path, excel_path)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
