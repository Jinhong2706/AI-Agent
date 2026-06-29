#!/usr/bin/env python3
"""
文档导入器 - 解析 Word/PDF 制度文档，智能分块并建立可搜索索引。

支持格式: .docx (Word), .pdf (PDF)
分块策略: 按章节 + 段落边界，保留重叠上下文
索引引擎: jieba 分词 + TF-IDF 权重

用法:
  python doc_importer.py --input employee_handbook.docx --category 员工手册 --name "公司员工手册V3.0"
  python doc_importer.py --list                          # 列出已导入文档
  python doc_importer.py --stats                         # 查看索引统计
"""

import argparse
import hashlib
import json
import os
import re
import sys
import time
from pathlib import Path
from datetime import datetime
from collections import defaultdict
import math

# 尝试导入 jieba 分词（临时重定向 stdout，屏蔽初始化打印）
import sys
import io
try:
    _saved_stdout = sys.stdout
    sys.stdout = io.StringIO()
    import jieba
    sys.stdout = _saved_stdout
    jieba.setLogLevel(50)
    HAS_JIEBA = True
except ImportError:
    sys.stdout = _saved_stdout if '_saved_stdout' in dir() else sys.__stdout__
    HAS_JIEBA = False
    print("[警告] jieba 未安装，将使用字符级分词（精度较低）。建议: pip install jieba")


# ============================================================
# 常量
# ============================================================

SKILL_DIR = Path(__file__).resolve().parent.parent
DATA_DIR = SKILL_DIR / "data"
REGISTRY_FILE = DATA_DIR / "docs_registry.json"
CHUNKS_DIR = DATA_DIR / "chunks"

# 分块参数
MAX_CHUNK_CHARS = 800       # 单块最大字符数
MIN_CHUNK_CHARS = 100       # 单块最小字符数（低于此合并到上一块）
OVERLAP_CHARS = 150         # 块与块之间的重叠字符数

# 文档分类
VALID_CATEGORIES = ["员工手册", "考勤制度", "福利政策", "报销制度", "薪酬制度", "绩效考核", "培训发展", "安全规范", "其他制度"]

# 中文标点 / 句式分隔符
SENTENCE_SPLITTERS = re.compile(r'[。！？；\n]{1,3}')
CHAPTER_PATTERNS = [
    re.compile(r'^第[一二三四五六七八九十百千\d]+章[\s　]+.{2,50}'),   # 第X章 标题
    re.compile(r'^第[一二三四五六七八九十百千\d]+节[\s　]+.{2,50}'),   # 第X节
    re.compile(r'^[一二三四五六七八九十]、.{2,50}'),                   # 一、标题
    re.compile(r'^\d+[\.\s、)]\s*.{2,50}'),                             # 1. 标题
    re.compile(r'^[（(][一二三四五六七八九十\d]+[）)]\s*.{0,50}'),     # (一) 标题
]


# ============================================================
# 文档解析
# ============================================================

def parse_docx(file_path: str) -> list[dict]:
    """解析 Word .docx 文件，返回带结构的段落列表"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")

    doc = Document(file_path)
    paragraphs = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # 推断段落样式层级
        style_name = para.style.name if para.style else ""
        level = _infer_level(style_name, text)

        paragraphs.append({
            "index": i,
            "text": text,
            "level": level,
            "style": style_name,
            "is_heading": level <= 2 or "Heading" in style_name or "标题" in style_name,
        })

    # 也尝试提取表格内容
    for ti, table in enumerate(doc.tables):
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        if rows_text:
            table_text = "\n".join(rows_text)
            if len(table_text) > 20:
                paragraphs.append({
                    "index": len(paragraphs),
                    "text": f"[表格] {table_text}",
                    "level": 3,
                    "style": "Table",
                    "is_heading": False,
                })

    return paragraphs


def parse_pdf(file_path: str) -> list[dict]:
    """解析 PDF 文件，优先使用 pdfplumber（更好处理表格），fallback PyPDF2"""
    text_content = ""

    # 尝试 pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n"
                # 提取表格
                for table in page.extract_tables():
                    rows_text = []
                    for row in table:
                        cells = [str(c).strip() if c else "" for c in row]
                        rows_text.append(" | ".join(cells))
                    if rows_text:
                        text_content += "[表格] " + "\n".join(rows_text) + "\n"
        if text_content.strip():
            return _raw_text_to_paragraphs(text_content)
    except Exception as e:
        print(f"[信息] pdfplumber 解析失败 ({e})，尝试 PyPDF2...")

    # Fallback: PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_content += page_text + "\n"
    except Exception as e:
        raise RuntimeError(f"PDF 解析失败: {e}")

    return _raw_text_to_paragraphs(text_content)


def _raw_text_to_paragraphs(text: str) -> list[dict]:
    """将原始文本按行分割为段落结构"""
    lines = text.strip().split("\n")
    paragraphs = []
    idx = 0
    for line in lines:
        line = line.strip()
        if not line:
            continue
        level = _infer_level("", line)
        paragraphs.append({
            "index": idx,
            "text": line,
            "level": level,
            "style": "",
            "is_heading": level <= 2 or _looks_like_heading(line),
        })
        idx += 1
    return paragraphs


def _infer_level(style_name: str, text: str) -> int:
    """推断段落层级：1=章标题, 2=节标题, 3=子标题, 4=正文"""
    if "Heading 1" in style_name or "标题 1" in style_name:
        return 1
    if "Heading 2" in style_name or "标题 2" in style_name:
        return 2

    for pattern in CHAPTER_PATTERNS:
        if pattern.match(text):
            if re.match(r'^第[一二三四五六七八九十百千\d]+章', text):
                return 1
            if re.match(r'^第[一二三四五六七八九十百千\d]+节', text):
                return 2
            return 3

    if _looks_like_heading(text):
        return 3

    return 4


def _looks_like_heading(text: str) -> bool:
    """启发式判断是否像标题"""
    if len(text) > 50:
        return False
    if text.endswith(("：", ":", "——")):
        return True
    if re.match(r'^[一二三四五六七八九十]+[、．.]', text):
        return True
    if re.match(r'^\d+[\.\s、)]', text) and len(text) < 40:
        return True
    return False


# ============================================================
# 智能分块
# ============================================================

def smart_chunk(paragraphs: list[dict], doc_name: str = "") -> list[dict]:
    """
    智能分块策略:
    1. 保持章节边界 — 新章节开始新块
    2. 按句号、换行等自然边界切分
    3. 块间保留 OVERLAP_CHARS 重叠上下文
    4. 表格内容尽量完整保留
    """
    chunks = []
    current_chunk = ""
    current_start = 0
    para_idx_in_chunk = 0
    chunk_meta = {"level": 4, "section": ""}

    for pi, para in enumerate(paragraphs):
        text = para["text"]
        level = para["level"]

        # 遇到章节标题 → 开始新块
        if level <= 2 and current_chunk.strip():
            chunks.append(_build_chunk(current_chunk, current_start, para_idx_in_chunk,
                                       chunk_meta["section"], chunk_meta["level"], doc_name))
            current_chunk = ""
            para_idx_in_chunk = 0

        # 更新当前节的标题
        if level <= 2:
            chunk_meta["section"] = text
            chunk_meta["level"] = level

        # 表格内容保留完整
        if text.startswith("[表格]"):
            if current_chunk.strip():
                chunks.append(_build_chunk(current_chunk, current_start, para_idx_in_chunk,
                                           chunk_meta["section"], chunk_meta["level"], doc_name))
            # 表格单独成块
            chunks.append(_build_chunk(text, pi, 1, chunk_meta["section"], chunk_meta["level"], doc_name))
            current_chunk = ""
            para_idx_in_chunk = 0
            current_start = pi + 1
            continue

        # 累积段落
        if not current_chunk:
            current_start = pi
            para_idx_in_chunk = 0

        current_chunk += text + "\n"
        para_idx_in_chunk += 1

        # 达到最大块大小时切分
        if len(current_chunk) >= MAX_CHUNK_CHARS:
            # 在句子边界切分
            split_pos = _find_sentence_boundary(current_chunk, MAX_CHUNK_CHARS - OVERLAP_CHARS)
            main_text = current_chunk[:split_pos]
            overlap_text = current_chunk[max(0, split_pos - OVERLAP_CHARS):]

            chunks.append(_build_chunk(main_text, current_start, para_idx_in_chunk,
                                       chunk_meta["section"], chunk_meta["level"], doc_name))

            # 下一块从重叠部分开始
            current_chunk = overlap_text
            current_start = pi - max(1, int(para_idx_in_chunk * 0.3))
            para_idx_in_chunk = max(1, int(para_idx_in_chunk * 0.3))

    # 处理最后一个块
    if current_chunk.strip():
        chunks.append(_build_chunk(current_chunk, current_start, para_idx_in_chunk,
                                   chunk_meta["section"], chunk_meta["level"], doc_name))

    # 合并过小的块
    chunks = _merge_small_chunks(chunks)

    return chunks


def _find_sentence_boundary(text: str, max_pos: int) -> int:
    """在 max_pos 附近寻找最近的句子边界"""
    if max_pos >= len(text):
        return len(text)

    # 在 max_pos 前后搜索句号、换行等边界
    search_window = min(100, len(text) - max_pos)
    for i in range(max_pos, min(max_pos + search_window, len(text))):
        if text[i] in "。！？\n":
            return i + 1
    for i in range(max_pos, max(0, max_pos - 50), -1):
        if text[i] in "。！？\n":
            return i + 1
    return max_pos


def _build_chunk(text: str, start_idx: int, para_count: int,
                 section: str, level: int, doc_name: str) -> dict:
    """构建标准块结构"""
    text = text.strip()
    chunk_id = hashlib.md5(text.encode("utf-8")).hexdigest()[:12]
    return {
        "chunk_id": chunk_id,
        "text": text,
        "char_count": len(text),
        "para_start": start_idx,
        "para_count": para_count,
        "section": section,
        "level": level,
        "doc_name": doc_name,
    }


def _merge_small_chunks(chunks: list[dict]) -> list[dict]:
    """合并过小的块到前一相邻块"""
    if len(chunks) <= 1:
        return chunks
    merged = []
    for chunk in chunks:
        if chunk["char_count"] < MIN_CHUNK_CHARS and merged:
            prev = merged[-1]
            prev["text"] = prev["text"] + "\n" + chunk["text"]
            prev["char_count"] = len(prev["text"])
            prev["para_count"] += chunk["para_count"]
        else:
            merged.append(chunk)
    return merged


# ============================================================
# 中文分词 & TF-IDF 索引
# ============================================================

def tokenize(text: str) -> list[str]:
    """中文分词 + 去停用词"""
    # 简单停用词表
    stop_words = {
        "的", "了", "在", "是", "我", "有", "和", "就", "不", "人", "都", "一",
        "一个", "上", "也", "很", "到", "说", "要", "去", "你", "会", "着",
        "没有", "看", "好", "自己", "这", "他", "她", "它", "们", "那", "些",
        "所", "被", "把", "让", "对", "从", "但", "而", "且", "与", "或",
        "可以", "需要", "应该", "必须", "可能", "已经", "将", "等", "其",
        "为", "以", "之", "及", "向", "由", "于", "中", "并", "如",
        "规定", "按照", "根据", "相关", "以下", "以上", "包括", "具体",
        "进行", "情况", "单位", "公司", "部门", "员工", "工作",
        "\n", "\r", "\t", " ", "  ",
    }

    if HAS_JIEBA:
        words = jieba.lcut(text)
    else:
        # 字符级 fallback: 按标点切分
        words = re.split(r'[，。！？；、\s\n（）()《》""\'\'【】：:]+', text)
        words = [w for w in words if len(w) >= 2]

    return [w for w in words if w not in stop_words and len(w) >= 2]


def build_tfidf_index(chunks: list[dict]) -> dict:
    """构建 TF-IDF 索引"""
    # 分词所有块
    chunk_tokens = []
    for chunk in chunks:
        tokens = tokenize(chunk["text"])
        chunk_tokens.append(tokens)

    # 计算 DF (document frequency)
    N = len(chunks)
    df = defaultdict(int)
    for tokens in chunk_tokens:
        unique_terms = set(tokens)
        for term in unique_terms:
            df[term] += 1

    # 构建每个块的 TF-IDF 向量
    vectors = []
    for tokens in chunk_tokens:
        tf = defaultdict(int)
        for t in tokens:
            tf[t] += 1
        # 归一化 TF
        max_tf = max(tf.values()) if tf else 1
        vec = {}
        for term, count in tf.items():
            idf = math.log((N + 1) / (df[term] + 1)) + 1
            vec[term] = (count / max_tf) * idf
        vectors.append(vec)

    # 构建倒排索引
    inverted_index = defaultdict(list)
    for ci, tokens in enumerate(chunk_tokens):
        for t in set(tokens):
            inverted_index[t].append(ci)

    return {
        "N": N,
        "chunks": chunks,
        "vectors": vectors,
        "inverted_index": dict(inverted_index),
        "df": dict(df),
        "built_at": datetime.now().isoformat(),
    }


# ============================================================
# 文档注册表管理
# ============================================================

def load_registry() -> dict:
    """加载文档注册表"""
    if REGISTRY_FILE.exists():
        with open(REGISTRY_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"documents": {}, "updated_at": ""}


def save_registry(registry: dict):
    """保存文档注册表"""
    registry["updated_at"] = datetime.now().isoformat()
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    with open(REGISTRY_FILE, "w", encoding="utf-8") as f:
        json.dump(registry, f, ensure_ascii=False, indent=2)


def import_document(file_path: str, category: str, name: str = "") -> dict:
    """
    导入文档主流程:
    1. 解析 → 2. 分块 → 3. 索引 → 4. 注册
    """
    file_path = os.path.abspath(file_path)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    # 检测格式
    ext = Path(file_path).suffix.lower()
    if ext not in (".docx", ".pdf"):
        raise ValueError(f"不支持的格式: {ext}，仅支持 .docx 和 .pdf")

    if category not in VALID_CATEGORIES:
        print(f"[警告] 分类 '{category}' 不在预设列表中，将使用自定义分类")

    # 自动生成文档名
    if not name:
        name = Path(file_path).stem

    # 解析文档
    print(f"[1/4] 解析文档: {file_path}")
    t0 = time.time()
    if ext == ".docx":
        paragraphs = parse_docx(file_path)
    else:
        paragraphs = parse_pdf(file_path)
    print(f"    解析完成: {len(paragraphs)} 个段落 ({time.time()-t0:.2f}s)")

    # 智能分块
    print(f"[2/4] 智能分块...")
    t0 = time.time()
    chunks = smart_chunk(paragraphs, name)
    print(f"    分块完成: {len(chunks)} 个块 ({time.time()-t0:.2f}s)")

    # 建立索引
    print(f"[3/4] 建立 TF-IDF 索引...")
    t0 = time.time()
    index = build_tfidf_index(chunks)

    # 保存索引到文件
    doc_id = hashlib.md5(f"{category}_{name}_{datetime.now().isoformat()}".encode()).hexdigest()[:16]
    index_file = CHUNKS_DIR / f"{doc_id}.json"
    CHUNKS_DIR.mkdir(parents=True, exist_ok=True)

    with open(index_file, "w", encoding="utf-8") as f:
        json.dump(index, f, ensure_ascii=False, indent=2)
    print(f"    索引完成: {index['N']} 个块, {len(index['df'])} 个词条 ({time.time()-t0:.2f}s)")

    # 注册文档
    print(f"[4/4] 注册文档...")
    registry = load_registry()
    file_stat = os.stat(file_path)

    doc_entry = {
        "doc_id": doc_id,
        "name": name,
        "category": category,
        "file_path": file_path,
        "file_type": ext,
        "file_size": file_stat.st_size,
        "file_mtime": datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
        "imported_at": datetime.now().isoformat(),
        "paragraphs_count": len(paragraphs),
        "chunks_count": len(chunks),
        "terms_count": len(index["df"]),
        "index_file": str(index_file),
    }

    # 检查是否同名覆盖
    existing_id = None
    for eid, entry in registry["documents"].items():
        if entry["name"] == name and entry["category"] == category:
            existing_id = eid
            break

    if existing_id:
        print(f"    [更新] 覆盖已有文档: {name} ({existing_id})")
        # 删除旧索引文件
        old_index = registry["documents"][existing_id].get("index_file", "")
        if old_index and os.path.exists(old_index):
            os.remove(old_index)

    registry["documents"][doc_id] = doc_entry
    save_registry(registry)

    # 返回摘要
    return {
        "doc_id": doc_id,
        "name": name,
        "category": category,
        "paragraphs": len(paragraphs),
        "chunks": len(chunks),
        "terms": len(index["df"]),
        "is_update": existing_id is not None,
    }


# ============================================================
# 公开查询接口
# ============================================================

def list_documents() -> list[dict]:
    """列出所有已导入文档"""
    registry = load_registry()
    docs = []
    for doc_id, entry in registry["documents"].items():
        docs.append({
            "doc_id": doc_id,
            "name": entry["name"],
            "category": entry["category"],
            "file_type": entry["file_type"],
            "file_size": entry["file_size"],
            "chunks": entry["chunks_count"],
            "imported_at": entry["imported_at"],
        })
    docs.sort(key=lambda d: d["imported_at"], reverse=True)
    return docs


def get_stats() -> dict:
    """获取索引统计"""
    registry = load_registry()
    total_chunks = sum(d["chunks_count"] for d in registry["documents"].values())
    total_terms = sum(d["terms_count"] for d in registry["documents"].values())
    categories = defaultdict(int)
    for d in registry["documents"].values():
        categories[d["category"]] += 1

    return {
        "documents_count": len(registry["documents"]),
        "total_chunks": total_chunks,
        "total_terms": total_terms,
        "categories": dict(categories),
        "last_updated": registry["updated_at"],
    }


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="员工手册文档导入器")
    parser.add_argument("--input", "-i", help="导入文档路径 (.docx / .pdf)")
    parser.add_argument("--name", "-n", default="", help="文档名称（默认使用文件名）")
    parser.add_argument("--category", "-c", default="其他制度",
                        help=f"文档分类: {', '.join(VALID_CATEGORIES)}")
    parser.add_argument("--list", "-l", action="store_true", help="列出已导入文档")
    parser.add_argument("--stats", "-s", action="store_true", help="查看索引统计")
    parser.add_argument("--output", "-o", default="", help="输出结果到 JSON 文件")

    args = parser.parse_args()

    result = {}

    if args.list:
        docs = list_documents()
        result = {"documents": docs, "count": len(docs)}
        if not docs:
            print("📭 暂无已导入文档")
        else:
            print(f"\n📚 已导入文档 ({len(docs)} 份):")
            print("-" * 70)
            for d in docs:
                size_kb = d["file_size"] / 1024
                print(f"  [{d['doc_id'][:8]}...] {d['name']}")
                print(f"    分类: {d['category']} | 格式: {d['file_type']} | 大小: {size_kb:.1f}KB | 块数: {d['chunks']}")
                print(f"    导入时间: {d['imported_at']}")
            print("-" * 70)

    elif args.stats:
        result = get_stats()
        print(f"\n📊 索引统计:")
        print(f"  文档总数: {result['documents_count']}")
        print(f"  总块数: {result['total_chunks']}")
        print(f"  总词条数: {result['total_terms']}")
        print(f"  分类分布: {result['categories']}")
        print(f"  最后更新: {result['last_updated']}")

    elif args.input:
        print(f"\n🚀 开始导入文档...\n")
        result = import_document(args.input, args.category, args.name)
        print(f"\n✅ 导入完成!")
        print(f"  文档ID: {result['doc_id']}")
        print(f"  文档名: {result['name']}")
        print(f"  分类: {result['category']}")
        print(f"  段落数: {result['paragraphs']}")
        print(f"  分块数: {result['chunks']}")
        print(f"  索引词条: {result['terms']}")
        if result["is_update"]:
            print(f"  ℹ️  已覆盖旧版本文档")

    else:
        parser.print_help()

    # 输出到文件
    if args.output and result:
        with open(args.output, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"\n📁 结果已保存到: {args.output}")

    # 返回 JSON 给调用方
    if not args.output:
        print(f"\n[JSON_RESULT] {json.dumps(result, ensure_ascii=False)}")


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding='utf-8') if hasattr(sys.stdout, 'reconfigure') else None
    main()
